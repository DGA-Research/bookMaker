
"""Local CLI for assembling briefing books from DOCX parts."""

import argparse
import platform
import sys
import tempfile
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SECTION_ORDER: Tuple[str, ...] = (
    "Top Hits",
    "Methodology",
    "Biographical",
    "Family/Personal Info",
    "Buisness Interests",
    "Race Review",
    "Campaign Finance",
    "Issues",
    "Appendicies",
    "Questionaires",
    "Scorecards",
    "Travel Discosureles",
    "Offical Office Disbursments",
)

FILE_NAME_MAP = {
    "Top Hits": "TOP HITS.docx",
    "Methodology": "METHODOLOGY.docx",
    "Biographical": "BIOGRAPHICAL.docx",
    "Family/Personal Info": "FAMILY PERSONAL INFO.docx",
    "Buisness Interests": "BUISNESS INTERESTS.docx",
    "Race Review": "RACE REVIEW.docx",
    "Campaign Finance": "CAMPAIGN FINANCE.docx",
    "Issues": "ISSUES.docx",
    "Appendicies": "APPENDICIES.docx",
    "Questionaires": "QUESTIONNAIRES.docx",
    "Scorecards": "SCORECARD.docx",
    "Travel Discosureles": "TRAVEL DISCLOSURES.docx",
    "Offical Office Disbursments": "OFFICIAL OFFICE DISBURSEMENTS.docx",
}

WD_PAGE_BREAK = 7
WD_STORY = 6
WD_HEADER_FOOTER_PRIMARY = 1
WD_ALIGN_PARAGRAPH_CENTER = 1


class LocalUpload:
    """Minimal file-like wrapper to mimic Streamlit uploads."""

    def __init__(self, path: Path) -> None:
        self.name = path.name
        self._bytes = path.read_bytes()

    def getvalue(self) -> bytes:
        return self._bytes


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Combine section DOCX files into a single briefing book.",
    )
    parser.add_argument(
        "--parts-dir",
        default="bookParts",
        help="Directory containing the section DOCX files (default: bookParts).",
    )
    parser.add_argument(
        "--output",
        default="combined_book.docx",
        help="Path for the generated DOCX (default: combined_book.docx).",
    )
    parser.add_argument(
        "--method",
        choices=["auto", "word", "python-docx"],
        default="auto",
        help=(
            "word: use Microsoft Word automation; python-docx: pure Python merge; "
            "auto: try word then fall back."
        ),
    )
    parser.add_argument(
        "--quiet",
        action="store_true",
        help="Suppress non-error output.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    book_parts_dir = Path(args.parts_dir).resolve()
    if not book_parts_dir.exists():
        sys.exit(f"bookParts directory not found: {book_parts_dir}")

    section_payloads = collect_section_payloads(book_parts_dir)
    if not section_payloads:
        sys.exit("No DOCX files discovered in the sections order.")

    filtered_sections: List[Tuple[str, List[LocalUpload]]] = []
    for section, paths in section_payloads:
        uploads = [LocalUpload(path) for path in paths]
        filtered_sections.append((section, uploads))

    method = resolve_method_choice(args.method)
    if method == "word" and not word_automation_available():
        print(
            "Microsoft Word automation unavailable; falling back to python-docx merge.",
            file=sys.stderr,
        )
        method = "python-docx"

    if method == "word":
        buffer = build_combined_document_with_word(filtered_sections)
    else:
        buffer = build_combined_document(filtered_sections)

    output_path = Path(args.output).resolve()
    output_path.write_bytes(buffer.getvalue())

    if not args.quiet:
        print(f"Combined document written to {output_path}")
        if method == "word":
            print("Open the document in Word to confirm the TOC and pagination look correct.")
        else:
            print("Open the document in Word and update fields (Ctrl+A, F9) to refresh the TOC.")


def resolve_method_choice(choice: str) -> str:
    if choice == "auto":
        return "word" if word_automation_available() else "python-docx"
    return choice


def collect_section_payloads(book_parts_dir: Path) -> List[Tuple[str, List[Path]]]:
    payloads: List[Tuple[str, List[Path]]] = []
    for section in SECTION_ORDER:
        paths = list(iter_section_files(book_parts_dir, section))
        if paths:
            payloads.append((section, paths))
        else:
            print(
                f"Warning: no DOCX files found for section '{section}' in {book_parts_dir}",
                file=sys.stderr,
            )
    return payloads


def iter_section_files(book_parts_dir: Path, section: str) -> Iterable[Path]:
    safe_dir = book_parts_dir / _safe_stem(section)
    if safe_dir.exists() and safe_dir.is_dir():
        yield from sorted(
            (p for p in safe_dir.glob("*.docx") if p.is_file()),
            key=lambda p: p.name.lower(),
        )
        return

    filename = FILE_NAME_MAP.get(section)
    if filename:
        candidate = book_parts_dir / filename
        if candidate.exists():
            yield candidate


def word_automation_available() -> bool:
    if platform.system().lower() != "windows":
        return False
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False
    return True


def build_combined_document(filtered_sections: List[Tuple[str, List]]) -> BytesIO:
    combined = Document()
    remove_initial_paragraph_if_empty(combined)
    section_style_name = ensure_section_style(combined)

    add_table_of_contents(combined, section_style_name)

    for section_name, files in filtered_sections:
        heading_text = section_name
        section_documents = []
        heading_captured = False

        for uploaded_file in files:
            file_bytes = uploaded_file.getvalue()
            source_doc = Document(BytesIO(file_bytes))
            if not heading_captured:
                heading_text = extract_section_heading_text(source_doc, section_name)
                heading_captured = True
            else:
                strip_leading_empty_paragraphs(source_doc)
            section_documents.append(source_doc)

        heading = combined.add_paragraph(heading_text, style=section_style_name)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.page_break_before = True

        for source_doc in section_documents:
            append_document_body(combined, source_doc)

    apply_footer_with_page_numbers(combined)

    output_buffer = BytesIO()
    combined.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer


def build_combined_document_with_word(filtered_sections: List[Tuple[str, List]]) -> BytesIO:
    if platform.system().lower() != "windows":
        raise RuntimeError("Microsoft Word automation is only supported on Windows.")

    try:
        import win32com.client as win32  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "pywin32 is required for Word-based composition. Install it with 'pip install pywin32'."
        ) from exc

    with tempfile.TemporaryDirectory() as tmpdir_str:
        tmpdir = Path(tmpdir_str)
        section_payloads: List[Tuple[str, List[Path]]] = []
        for section_name, files in filtered_sections:
            heading_text = section_name
            file_paths: List[Path] = []
            heading_captured = False
            for index, uploaded_file in enumerate(files):
                suffix = Path(getattr(uploaded_file, "name", "") or "document.docx").suffix or ".docx"
                safe_name = _safe_stem(section_name)
                file_path = tmpdir / f"{safe_name}_{index}{suffix}"
                file_path.write_bytes(uploaded_file.getvalue())

                doc = Document(file_path)
                if not heading_captured:
                    heading_text = extract_section_heading_text(doc, section_name)
                    heading_captured = True
                else:
                    strip_leading_empty_paragraphs(doc)
                doc.save(file_path)

                file_paths.append(file_path)
            section_payloads.append((heading_text, file_paths))

        output_path = tmpdir / "combined.docx"
        compose_sections_with_word(section_payloads, output_path)
        return BytesIO(output_path.read_bytes())


def compose_sections_with_word(section_payloads: List[Tuple[str, List[Path]]], output_path: Path) -> None:
    import win32com.client as win32  # type: ignore

    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = None
    try:
        doc = word.Documents.Add()
        selection = word.Selection
        insert_table_of_contents_word(doc, selection)

        for index, (heading_text, files) in enumerate(section_payloads):
            if index > 0:
                selection.InsertBreak(WD_PAGE_BREAK)
            selection.Style = "Heading 1"
            selection.TypeText(heading_text)
            selection.TypeParagraph()
            for file_path in files:
                selection.InsertFile(str(file_path))
                selection.EndKey(Unit=WD_STORY)

        if doc.TablesOfContents.Count:
            doc.TablesOfContents(1).Update()
        apply_footer_with_page_numbers_word(doc)
        doc.SaveAs(str(output_path))
    finally:
        if doc is not None:
            doc.Close(SaveChanges=False)
        word.Quit()


def insert_table_of_contents_word(doc, selection) -> None:
    selection.Style = "Title"
    selection.TypeText("Table of Contents")
    selection.TypeParagraph()
    doc.TablesOfContents.Add(
        Range=selection.Range,
        UseHeadingStyles=True,
        UpperHeadingLevel=1,
        LowerHeadingLevel=3,
        UseHyperlinks=True,
        IncludePageNumbers=True,
        RightAlignPageNumbers=True,
    )
    selection.EndKey(Unit=WD_STORY)
    selection.TypeParagraph()
    selection.InsertBreak(WD_PAGE_BREAK)
    selection.EndKey(Unit=WD_STORY)


def apply_footer_with_page_numbers_word(doc) -> None:
    for section in doc.Sections:
        footer = section.Footers(WD_HEADER_FOOTER_PRIMARY)
        footer.Range.Text = ""
        footer.PageNumbers.RestartNumberingAtSection = False
        footer.PageNumbers.Add(PageNumberAlignment=WD_ALIGN_PARAGRAPH_CENTER)
        footer.Range.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER


def _safe_stem(value: str) -> str:
    stem = ''.join(ch if ch.isalnum() else '_' for ch in value).strip('_')
    return stem or "section"


def ensure_section_style(document: Document) -> str:
    style_name = "BookMaker Section"
    try:
        document.styles[style_name]
        return style_name
    except KeyError:
        pass

    section_style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = document.styles["Heading 1"]
    section_style.quick_style = True
    return style_name


def add_table_of_contents(document: Document, section_style_name: str) -> None:
    document.add_paragraph("Table of Contents", style="Title")
    toc_paragraph = document.add_paragraph()
    create_field_run(
        toc_paragraph,
        f'TOC \\h \\z \\t "{section_style_name},1"',
        "Update this table in Word to populate the entries.",
    )


def apply_footer_with_page_numbers(document: Document) -> None:
    for section in document.sections:
        section.footer.is_linked_to_previous = False
        if section.footer.paragraphs:
            footer_paragraph = section.footer.paragraphs[0]
            clear_paragraph(footer_paragraph)
        else:
            footer_paragraph = section.footer.add_paragraph()

        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.add_run("Page ")
        create_field_run(footer_paragraph, "PAGE", "1")
        footer_paragraph.add_run(" of ")
        create_field_run(footer_paragraph, "NUMPAGES", "1")


def append_document_body(target: Document, source: Document) -> None:
    for element in list(source.element.body):
        if element.tag == qn("w:sectPr"):
            continue
        target.element.body.append(deepcopy(element))


def paragraph_has_visible_content(paragraph) -> bool:
    if paragraph.text.strip():
        return True
    element = paragraph._element
    return bool(element.xpath('.//w:drawing') or element.xpath('.//w:pict'))


def extract_section_heading_text(document: Document, fallback: str) -> str:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip():
            heading_text = paragraph.text.strip()
            remove_paragraph(paragraph)
            return heading_text
        if paragraph_has_visible_content(paragraph):
            continue
        remove_paragraph(paragraph)
    return fallback


def strip_leading_empty_paragraphs(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() or paragraph_has_visible_content(paragraph):
            break
        remove_paragraph(paragraph)


def create_field_run(paragraph, field_code: str, default_text: str = ""):
    run = paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(field_begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    run._r.append(instruction)

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(field_separate)

    text = OxmlElement("w:t")
    text.text = default_text
    run._r.append(text)

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_end)

    return run


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def remove_initial_paragraph_if_empty(document: Document) -> None:
    if document.paragraphs and not document.paragraphs[0].text:
        p = document.paragraphs[0]._p
        p.getparent().remove(p)


if __name__ == "__main__":
    main()
